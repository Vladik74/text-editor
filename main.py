import ast
import codecs
import inspect
import os
import time
import tkinter as tk
import unicodedata
import plugins.base_plugin
from tkinter import filedialog as fd
from tkinter import font, Text, messagebox
from tkinter.ttk import Combobox
from docx import Document
from importlib import import_module


class Model:
    def __init__(self, filename):
        self.filename = filename

    @staticmethod
    def write_tags(file, text_info, size):
        tags = text_info.dump("1.0", tk.END)
        tags_to_str = str(tags).strip('[]')
        file.write('`````')
        file.write(f'({size})')
        file.write(tags_to_str)
        print(tags_to_str)

    def open_file(self, text_info, fonts, root):
        self.filename = fd.askopenfilename(filetypes=[("Text files", "*.txt")],
                                           initialdir='/')
        with codecs.open(self.filename, errors='ignore') as f:
            file_size_in_bytes = os.stat(f.name).st_size
            if file_size_in_bytes >= 52428800:
                tk.messagebox.showerror("Exceeding file size",
                                        "Your file is too big for editor. "
                                        "There are may be problems with "
                                        "opening the file.")
                root.destroy()
            short_name = f.name.split('/')[-1]
            r_text_file = f.read()
        tags_pos = r_text_file.find('`````')
        if tags_pos != 1 and tags_pos != -1:
            text = r_text_file[:tags_pos]
        else:
            text = r_text_file
        text_info.delete(1.0, tk.END)
        text_info.insert(tk.END, text)
        if tags_pos != 1 and tags_pos != -1:
            tags = r_text_file[tags_pos + 9:]
            size = int(r_text_file[tags_pos + 6:tags_pos + 8])
            self.get_format_from_tags(tags, text_info, fonts)
        else:
            size = 14

        return short_name, size

    @staticmethod
    def create_file(text_info):
        text_info.delete(1.0, tk.END)
        return "New"

    def save_as_file(self, text_info, size):
        file_types = [("Text File", "*.txt"), ("All files", "*.*")]
        self.filename = fd.asksaveasfilename(confirmoverwrite=True,
                                             defaultextension=file_types,
                                             filetypes=file_types)
        if self.filename is None:
            return
        content = str(text_info.get(1.0, tk.END))
        save_file = open(self.filename, 'w+')
        normalized_content = unicodedata.normalize('NFC', content)
        save_file.write(normalized_content)
        time.sleep(3)
        self.write_tags(save_file, text_info, size)
        save_file.close()
        short_name = self.filename.split('/')[-1]

        return short_name

    def save(self, text_info, size):
        print(self.filename)
        if self.filename == "":
            self.save_as_file(text_info, size)
        else:
            content = str(text_info.get(1.0, tk.END))
            normalized_content = unicodedata.normalize('NFC', content)
            save_file = open(self.filename, 'w+')
            save_file.write(normalized_content)
            self.write_tags(save_file, text_info, size)
            save_file.close()
            short_name = self.filename.split('/')[-1]

            return short_name

    @staticmethod
    def export_to_docx(text_info):
        doc = Document()
        strings = [text_info.get('1.0', 'end-1c')]
        for string in strings:
            run = doc.add_paragraph(string).add_run()
            fontt = run.font
            fontt.italic = True
        doc.save("C:/Users/vlad0/OneDrive/Документы/tkdoc.docx")

    @staticmethod
    def get_format_from_tags(tags: str, text_info, fonts):
        tags_list = list(ast.literal_eval(tags))
        current_font = None
        start_symbol = None
        for t in tags_list:
            print(t)
            if t[0] == 'tagon' and t[1] in fonts:
                current_font = t[1]
                start_symbol = t[2]
                continue
            if t[0] == 'tagoff' and t[1] in fonts:
                text_info.tag_configure(current_font, font=fonts[current_font])
                text_info.tag_add(current_font, start_symbol, t[2])


class Controller:
    def __init__(self, model, view):
        self.model = model
        self.view = view

    def open_file(self):
        name, size = self.model.open_file(self.view.text_info, self.view.fonts,
                                          self.view.root)
        self.change_fonts(self.view.fonts, size)
        self.view.root.title(name)

    def create_file(self):
        name = self.model.create_file(self.view.text_info)
        self.view.root.title(name)

    def save_as_file(self):
        name = self.model.save_as_file(self.view.text_info, self.view.font_size)
        self.view.root.title(name)

    def save(self):
        name = self.model.save(self.view.text_info, self.view.font_size)
        self.view.root.title(name)

    def close_file(self):
        answer = tk.messagebox.askquestion("Save or close",
                                           "Do you want to save changes?")
        if answer == "yes":
            self.save()
        self.view.root.destroy()

    def export_to_docx(self):
        self.model.export_to_docx(self.view.text_info)

    def get_bold_text(self):
        self.view.text_info.tag_configure("bold", font=self.view.fonts["bold"])
        current_tags = self.view.text_info.tag_names("sel.first")
        if "bold" in current_tags:
            self.view.text_info.tag_remove("bold", "sel.first", "sel.last")
        else:
            self.view.text_info.tag_add("bold", "sel.first", "sel.last")

    def get_italic_text(self):
        self.view.text_info.tag_configure("italic",
                                          font=self.view.fonts["italic"])
        current_tags = self.view.text_info.tag_names("sel.first")
        if "italic" in current_tags:
            self.view.text_info.tag_remove("italic", "sel.first", "sel.last")
        else:
            self.view.text_info.tag_add("italic", "sel.first", "sel.last")

    def get_underline_text(self):
        self.view.text_info.tag_configure("underline",
                                          font=self.view.fonts["underline"])
        current_tags = self.view.text_info.tag_names("sel.first")
        if "underline" in current_tags:
            self.view.text_info.tag_remove("underline", "sel.first", "sel.last")
        else:
            self.view.text_info.tag_add("underline", "sel.first", "sel.last")

    def get_overstrike_text(self):
        self.view.text_info.tag_configure("overstrike",
                                          font=self.view.fonts["overstrike"])
        current_tags = self.view.text_info.tag_names("sel.first")
        if "overstrike" in current_tags:
            self.view.text_info.tag_remove("overstrike", "sel.first",
                                           "sel.last")
        else:
            self.view.text_info.tag_add("overstrike", "sel.first", "sel.last")

    def change_size(self, event):
        selected_size = int(event.widget.get())
        self.change_fonts(self.view.fonts, selected_size)

    def change_fonts(self, fonts: dict, selected_size):
        self.view.text_info.config(font=('Helvetica', selected_size))
        index = self.view.sizes.index(selected_size)
        self.view.sizes_combobox.current(index)
        self.view.font_size = selected_size
        fonts.update({"default": font.Font(size=selected_size),
                      "bold": font.Font(size=selected_size, weight='bold'),
                      "italic": font.Font(size=selected_size, slant='italic'),
                      "underline": font.Font(size=selected_size, underline=1),
                      "overstrike": font.Font(size=selected_size,
                                              overstrike=1)})
        self.view.text_info.tag_configure("bold", font=fonts["bold"])
        self.view.text_info.tag_configure("italic", font=fonts["italic"])
        self.view.text_info.tag_configure("underline", font=fonts["underline"])
        self.view.text_info.tag_configure("overstrike",
                                          font=fonts["overstrike"])


class View:
    def __init__(self):
        self.root = tk.Tk()
        self.scrollbar = tk.Scrollbar(self.root)
        self.root.geometry("1200x2400")
        self.root.minsize(400, 400)
        self.root.maxsize(1600, 3200)
        self.root.title("New")
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.fontbar_frame = tk.Frame(self.root)
        self.fontbar_frame.pack(fill=tk.X)
        self.plugins = []

        self.font_size = 12
        self.fonts = {
            "default": font.Font(size=self.font_size),
            "bold": font.Font(size=self.font_size, weight='bold'),
            "italic": font.Font(size=self.font_size, slant='italic'),
            "underline": font.Font(size=self.font_size, underline=1),
            "overstrike": font.Font(size=self.font_size, overstrike=1)
        }
        #  region create_buttons
        self.bold_button_font = font.Font(size=10, weight='bold')
        self.bold_button = tk.Button(self.fontbar_frame, text="Bold",
                                     font=self.bold_button_font)
        self.bold_button.pack(side=tk.LEFT)

        self.italic_button_font = font.Font(size=10, slant='italic')
        self.italic_button = tk.Button(self.fontbar_frame, text="Italic",
                                       font=self.italic_button_font)
        self.italic_button.pack(side=tk.LEFT)

        self.underline_button_font = font.Font(size=10, underline=1)
        self.underline_button = tk.Button(self.fontbar_frame, text="Underline",
                                          font=self.underline_button_font)
        self.underline_button.pack(side=tk.LEFT)

        self.overstrike_button_font = font.Font(size=10, overstrike=1)
        self.overstrike_button = tk.Button(self.fontbar_frame, text="Strikeout",
                                           font=self.overstrike_button_font)
        self.overstrike_button.pack(side=tk.LEFT)
        # endregion
        self.text_info: Text = tk.Text(self.root,
                                       yscrollcommand=self.scrollbar.set,
                                       font=self.fonts["default"], undo=True)
        self.text_info.pack(fill=tk.BOTH)

        self.scrollbar.config(command=self.text_info.yview)
        # region undo-redo

        self.undo_icon = tk.PhotoImage(file="icons/undo_icon.png").subsample(9,
                                                                             9)
        self.undo_button = tk.Button(self.fontbar_frame, image=self.undo_icon,
                                     command=self.text_info.edit_undo)
        self.redo_icon = tk.PhotoImage(file="icons/redo_icon.png").subsample(26,
                                                                             26)
        self.redo_button = tk.Button(self.fontbar_frame, image=self.redo_icon,
                                     command=self.text_info.edit_redo)

        self.undo_button.pack(side=tk.LEFT)
        self.redo_button.pack(side=tk.LEFT)

        # endregion

        self.sizes = [10, 12, 14, 16, 18, 20, 22, 24, 26, 28, 30, 32]
        self.sizes_combobox = Combobox(self.fontbar_frame, values=self.sizes,
                                       state="readonly")
        self.sizes_combobox.current(2)
        self.sizes_combobox.pack(side=tk.LEFT)

        self.menu = tk.Menu(self.root)
        self.file_menu = tk.Menu(self.menu, tearoff=0)

    def install_plugins(self):
        for plugin in os.listdir("plugins"):
            if plugin.endswith('.py'):
                plugin_name = plugin[:-3]
                if plugin_name != "base_plugin" and plugin_name != "__init__":
                    self.plugins.append("plugins" + "." + plugin_name)
        for plugin in self.plugins:
            plugin_obj = import_module(plugin)
            for e in dir(plugin_obj):
                obj = getattr(plugin_obj, e)
                if inspect.isclass(obj):
                    if issubclass(obj, plugins.base_plugin.BasePlugin):
                        plugin_installer = obj(self)
                        plugin_installer.install()

    def set_controller(self, editor_controller):
        self.file_menu.add_command(label="New",
                                   command=editor_controller.create_file)
        self.file_menu.add_command(label="Open",
                                   command=editor_controller.open_file)
        self.file_menu.add_command(label="Save", command=editor_controller.save)
        self.file_menu.add_command(label="Save as",
                                   command=editor_controller.save_as_file)
        self.file_menu.add_command(label="Close",
                                   command=editor_controller.close_file)
        self.file_menu.add_command(label="To docx",
                                   command=editor_controller.export_to_docx)
        self.menu.add_cascade(label="File", menu=self.file_menu)
        self.root.config(menu=self.menu)
        self.sizes_combobox.bind("<<ComboboxSelected>>",
                                 editor_controller.change_size)
        self.bold_button.configure(command=editor_controller.get_bold_text)
        self.italic_button.configure(command=editor_controller.get_italic_text)
        self.underline_button.configure(command=editor_controller.
                                        get_underline_text)
        self.overstrike_button.configure(command=editor_controller
                                         .get_overstrike_text)

        self.install_plugins()

        self.root.mainloop()


if __name__ == '__main__':
    model = Model("")
    view = View()
    controller = Controller(model, view)
    view.set_controller(controller)
