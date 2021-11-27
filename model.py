import ast
import codecs
import os
import time
import tkinter as tk
import unicodedata
from tkinter import filedialog as fd
from tkinter import font, messagebox
from docx import Document


class Model:
    def __init__(self, filename):
        self.filename = filename

    @staticmethod
    def write_tags(file, text_info, size):
        tags = text_info.dump("1.0", tk.END)
        tags_to_str = str(tags).strip('[]')
        file.write('`````')
        file.write(f'({size})')
        file.write(tags_to_str)
        # print(tags_to_str)

    def open_file(self, text_info, fonts, root):
        self.filename = fd.askopenfilename(filetypes=[("Text files", "*.txt")],
                                           initialdir='/')
        with codecs.open(self.filename, errors='ignore') as f:
            file_size_in_bytes = os.stat(f.name).st_size
            if file_size_in_bytes >= 52428800:
                tk.messagebox.showerror("Exceeding file size",
                                        "Your file is too big for editor. "
                                        "There are may be problems with "
                                        "opening the file.")
                root.destroy()
            short_name = f.name.split('/')[-1]
            r_text_file = f.read()
        tags_pos = r_text_file.find('`````')
        if tags_pos != 1 and tags_pos != -1:
            text = r_text_file[:tags_pos]
        else:
            text = r_text_file
        text_info.delete(1.0, tk.END)
        text_info.insert(tk.END, text)
        if tags_pos != 1 and tags_pos != -1:
            tags = r_text_file[tags_pos + 9:]
            size = int(r_text_file[tags_pos + 6:tags_pos + 8])
            self.get_format_from_tags(tags, text_info, fonts)
        else:
            size = 14

        return short_name, size

    @staticmethod
    def create_file(text_info):
        text_info.delete(1.0, tk.END)
        return "New"

    def save_as_file(self, text_info, size):
        file_types = [("Text File", "*.txt"), ("All files", "*.*")]
        self.filename = fd.asksaveasfilename(confirmoverwrite=True,
                                             defaultextension=file_types,
                                             filetypes=file_types)
        if self.filename is None:
            return
        content = str(text_info.get(1.0, tk.END))
        save_file = open(self.filename, 'w+')
        normalized_content = unicodedata.normalize('NFC', content)
        save_file.write(normalized_content)
        time.sleep(3)
        self.write_tags(save_file, text_info, size)
        save_file.close()
        short_name = self.filename.split('/')[-1]

        return short_name

    def save(self, text_info, size):
        if self.filename == "":
            self.save_as_file(text_info, size)
        else:
            content = str(text_info.get(1.0, tk.END))
            normalized_content = unicodedata.normalize('NFC', content)
            save_file = open(self.filename, 'w+')
            save_file.write(normalized_content)
            self.write_tags(save_file, text_info, size)
            save_file.close()
            short_name = self.filename.split('/')[-1]

            return short_name

    @staticmethod
    def export_to_docx(text_info):
        doc = Document()
        strings = [text_info.get('1.0', 'end-1c')]
        for string in strings:
            run = doc.add_paragraph(string).add_run()
            fontt = run.font
            fontt.italic = True
        doc.save("C:/Users/vlad0/OneDrive/Документы/tkdoc.docx")

    @staticmethod
    def get_format_from_tags(tags: str, text_info, fonts):
        tags_list = list(ast.literal_eval(tags))
        current_font = None
        start_symbol = None
        for t in tags_list:
            # print(t)
            if t[0] == 'tagon' and t[1] in fonts:
                current_font = t[1]
                start_symbol = t[2]
                continue
            if t[0] == 'tagoff' and t[1] in fonts:
                text_info.tag_configure(current_font, font=fonts[current_font])
                text_info.tag_add(current_font, start_symbol, t[2])
